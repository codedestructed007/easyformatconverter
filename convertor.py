import html
import os
from io import BytesIO, StringIO


import streamlit as st
import win32com.client.dynamic as win32
from streamlit_option_menu import option_menu

# -------------------------------------------------------
st.set_page_config(
    page_title="File Converter",
    page_icon=":file_folder:"
)


# About page starts---------------------------------------
def about_page() :
    st.title("About File Converter")
    st.markdown("---")

    st.header("Introduction")
    st.write("Welcome to File Converter, a powerful and user-friendly application developed by a single developer "
             "that allows you to convert files from one format to another.")

    st.header("Key Features")
    st.markdown("- Supports various file formats, including Word, CSV, PDF, TXT, and HTML.")
    st.markdown("- Simple and intuitive interface for easy file uploading and conversion.")
    st.markdown("- Fast and reliable conversion process, ensuring accurate results.")

    st.header("Conversion Options")
    st.markdown("Choose from a range of conversion options:")
    st.markdown("- Convert text files to Word, CSV, or PDF.")
    st.markdown("- Convert DOCX files to TXT, CSV, or PDF.")
    st.markdown("- Convert CSV files to TXT, Word, or PDF.")
    st.markdown("- Convert HTML files to TXT, CSV, Word, or PDF.")

    st.header("Future Updates")
    st.write("We are continuously working to enhance File Converter with more features:")
    st.markdown("- Audio file conversion support will be added in the upcoming updates.")
    st.markdown("- Video file conversion support will be added in the future versions.")
    st.markdown("- Complex document conversion, such as Markdown (.md) files, will be supported soon.")

    st.header("User Experience")
    st.write("Our application offers a seamless user experience, with features designed to enhance usability:")
    st.markdown("- Drag and drop files for quick and convenient upload.")
    st.markdown("- Real-time conversion previews for instant feedback.")
    st.markdown("- Progress indicators to track the conversion process.")

    st.header("Privacy and Security")
    st.write("We prioritize the privacy and security of your files:")
    st.markdown("- All uploaded files are handled with strict confidentiality.")
    st.markdown("- We take measures to protect user data and ensure its security.")

    st.header("Instructions")
    st.write("Here's how to use File Converter:")
    st.markdown("1. Upload your file by clicking the 'Upload a file' button.")
    st.markdown("2. Select the desired conversion format from the 'Convert into' dropdown.")
    st.markdown("3. Click the 'Convert' button(only in some cases) to start the conversion process.")
    st.markdown("4. Wait for the conversion to complete.")
    st.markdown(
        "5. Once the conversion is finished, click the 'Download Converted File' button to save the converted file.")

    st.header("Support and Feedback")
    st.write("We value your feedback and are here to assist you:")
    st.markdown("- For any questions or issues, reach out to our support team at codexistslonglastingnotfog@gmail.com.")
    st.markdown("- We appreciate your feedback and suggestions for improving our application.")


# About page ends------------------------------------

# Help/FAQ page starts---------------------------------------
def help_faq_page() :
    st.title("Help/FAQ")

    st.header("Frequently Asked Questions")

    st.markdown("Here are some frequently asked questions about File Converter:")

    st.markdown("Q: What file formats are supported by File Converter?")
    st.markdown("A: File Converter supports various file formats, including Word (DOCX), CSV, PDF, TXT, and HTML.")

    st.markdown("Q: How can I convert a file?")
    st.markdown("A: To convert a file, follow these steps:")
    st.markdown("1. Click the 'Upload a file' button to select the file you want to convert.")
    st.markdown("2. Choose the desired conversion format from the 'Convert into' dropdown.")
    st.markdown("3. Click the 'Convert' button to start the conversion process.")
    st.markdown("4. Wait for the conversion to complete.")
    st.markdown(
        "5. Once the conversion is finished, click the 'Download Converted File' button to save the converted file.")

    st.markdown("Q: Can I convert multiple files at once?")
    st.markdown(
        "A: Currently, File Converter supports converting one file at a time. You can convert multiple files by "
        "repeating the conversion process for each file.")

    st.markdown("Q: Is my data secure?")
    st.markdown(
        "A: Yes, we prioritize the privacy and security of your files. All uploaded files are handled with strict "
        "confidentiality, and we take measures to protect user data and ensure its security.")

    st.markdown("Q: How long does the conversion process take?")
    st.markdown(
        "A: The time taken for conversion depends on the size and complexity of the file. Larger files or files with "
        "complex formatting may take longer to convert.")

    st.markdown("Q: What if I encounter any issues or have additional questions?")
    st.markdown(
        "A: If you encounter any issues or have additional questions, please reach out to our support team at "
        "codexistslonglastingnotfog@gmail.com. We are here to assist you.")

    st.header("Additional Help")
    st.markdown(
        "If you need further assistance or have any other questions, feel free to contact us. We are always happy to "
        "help.")

    st.header("Feedback")
    st.markdown(
        "We value your feedback and suggestions for improving our application. If you have any ideas or feature "
        "requests, please let us know.")


# Help/FAQ page ends -----------------------------------------


# contact page starts here------------------------------------
def contact_page() :
    st.title("Contact")

    st.header("Contact Information")
    st.write("You can reach out to us through the following channels:")

    st.markdown("- Email: [codexistslonglastingnotfog@gmail.com](mailto:codexistslonglastingnotfog@gmail.com)")
    st.markdown("- Twitter: [@SatyamS67442021](https://twitter.com/SatyamS67442021?t=Lm-I23EPPxWfGRd1jMyKXg&s=09)")
    st.markdown("- LinkedIn: [Satyam Sharma](https://www.linkedin.com/in/satyamsharma61541425b)")

    st.header("Thank You")
    st.write("Thank you for your interest in our application. We look forward to hearing from you!")


# contact page ends herer-------------------------------------


# feedback page starts here---------------------------------
def feedback_page() :
    st.title("Feedback")

    st.header("We Value Your Feedback")
    st.write(
        "We appreciate your feedback and suggestions for improving our application. Your feedback helps us understand "
        "your needs better and allows us to enhance your user experience.")

    st.header("How to Provide Feedback")
    st.write("There are several ways you can provide feedback:")
    st.markdown(
        "- Send us an email at [codexistslonglastingnotfog@gmail.com](mailto:codexistslonglastingnotfog@gmail.com)")
    st.markdown("- Reach out to us through our social media channels (Twitter, LinkedIn)")
    st.markdown("- Fill out the feedback form on our website")

    st.header("What We're Interested In")
    st.write("We value feedback related to the following aspects:")
    st.markdown("- User experience and interface design suggestions")
    st.markdown("- Feature requests and additional functionality you'd like to see")
    st.markdown("- Bug reports and technical issues")
    st.markdown("- Performance improvements")

    st.header("Stay Updated")
    st.write("To stay updated with the latest news and updates, follow us on social media:")
    st.markdown("[LinkedIn]([Satyam Sharma](https://www.linkedin.com/in/satyamsharma61541425b)")

    st.header("Thank You")
    st.write(
        "Thank you for taking the time to provide us with your feedback. We value your input and continuously strive "
        "to improve our application based on your suggestions.")


# feedback page ends here-----------------------------


with st.sidebar :
    selected = option_menu(
        menu_title='Main Menu',
        options=['Home', 'About', 'Help/FAQ', 'Contact', 'Feedback'],
        icons=['house', 'info-circle', 'question-circle', 'envelope', 'exclamation-circle'],
        menu_icon='volume',  # optional
        default_index=0,
        styles={

        }
    )


# if selected == 'Home':


def home_page() :
    import time
    def text_to_word(uploaded_file) :
        output_file = BytesIO()
        text = uploaded_file.read().decode("utf-8")

        document = Document()
        document.add_paragraph(text)
        document.save(output_file)
        output_file.seek(0)
        # Convert button
        left, center, right = st.columns(3)
        with left :

            if st.button("Convert") :
                with right :
                    with st.spinner(f"Converting into Word...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            def get_key_by_value(dictionary, value) :
                                for key, val in dictionary.items() :
                                    if val == value :
                                        return key
                                # If the value is not found, you can return None or raise an exception.
                                return None

                            st.success(f"File converted to {selected_option} successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.{get_key_by_value(options, selected_option)}"

                            st.download_button("Download Converted File", data=output_file, file_name=new_file_name)

                        # st.download_button("Download Converted File", data=output_file, file_name=f"{uploaded_file.name[:-4]}.docx")

    def text_to_csv(uploaded_file) :
        output_file = BytesIO()

        text = uploaded_file.read().decode("utf-8")

        rows = text.split('\n')

        with StringIO() as file :
            csv_writer = csv.writer(file)
            for row in rows :
                csv_writer.writerow(row.split(','))

            output_file.write(file.getvalue().encode())

        output_file.seek(0)
        left, center, right = st.columns(3)
        with left :
            if st.button("Convert") :
                with right :
                    with st.spinner(f"Converting into Word...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            def get_key_by_value(dictionary, value) :
                                for key, val in dictionary.items() :
                                    if val == value :
                                        return key
                                # If the value is not found, you can return None or raise an exception.
                                return None

                            st.success(f"File converted to {selected_option} successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.{get_key_by_value(options, selected_option)}"

                            st.download_button("Download Converted File", data=output_file, file_name=new_file_name)
            # todo-fixme This has been done

    def text_to_pdf(uploaded_file) :
        text = uploaded_file.read().decode("utf-8")

        # Create a temporary file to save the PDF
        temp_dir = tempfile.mkdtemp()
        temp_file_path = f"{temp_dir}/converted.pdf"

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        for line in text.split('\n') :
            pdf.cell(0, 10, txt=line, ln=True, align='L')

        # Save the PDF to the temporary file
        pdf.output(temp_file_path)

        # Read the temporary PDF file as bytes
        with open(temp_file_path, 'rb') as temp_file :
            pdf_bytes = temp_file.read()

        left, center, right = st.columns(3)
        with left :
            if st.button("Convert") :
                with right :
                    with st.spinner(f"Converting into Word...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            def get_key_by_value(dictionary, value) :
                                for key, val in dictionary.items() :
                                    if val == value :
                                        return key
                                # If the value is not found, you can return None or raise an exception.
                                return None

                            st.success(f"File converted to {selected_option} successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.{get_key_by_value(options, selected_option)}"

                            st.download_button("Download Converted File", data=pdf_bytes, file_name=new_file_name)
            # todo-fixme checked

    def text_to_html(uploaded_file) :
        # Create a BytesIO object to store the converted HTML
        output_html = io.BytesIO()

        # Read the text file
        text_content = uploaded_file.read().decode()

        # Convert text to HTML
        html_content = f"<pre>{html.escape(text_content)}</pre>"

        # Write the HTML content to the BytesIO object
        output_html.write(html_content.encode())
        output_html.seek(0)

        left, center, right = st.columns(3)
        with left :
            if st.button("Convert") :
                with right :
                    with st.spinner("Converting to HTML...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            st.success("File converted to HTML successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.html"

                            st.download_button("Download Converted File", data=output_html, file_name=new_file_name)

                            # todo-fixme done

    # text to others ends ------------------------

    # docx to others begins ------------------------

    def docx_to_text(uploaded_file) :
        output_file = BytesIO()

        document = Document(uploaded_file)

        text = ""
        for paragraph in document.paragraphs :
            text += paragraph.text + "\n"

        output_file.write(text.encode())

        output_file.seek(0)

        left, center, right = st.columns(3)
        with left :

            if st.button("Convert") :
                with right :
                    with st.spinner(f"Converting into Word...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            def get_key_by_value(dictionary, value) :
                                for key, val in dictionary.items() :
                                    if val == value :
                                        return key
                                # If the value is not found, you can return None or raise an exception.
                                return None

                            st.success(f"File converted to {selected_option} successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.{get_key_by_value(options, selected_option)}"

                            st.download_button("Download Converted File", data=output_file, file_name=new_file_name)

    def docx_to_csv(uploaded_file) :
        output_file = BytesIO()

        document = Document(uploaded_file)

        rows = []
        for paragraph in document.paragraphs :
            rows.append(paragraph.text.split(','))

        with StringIO() as file :
            csv_writer = csv.writer(file)
            csv_writer.writerows(rows)

            output_file.write(file.getvalue().encode())

        output_file.seek(0)

        left, center, right = st.columns(3)
        with left :

            if st.button("Convert") :
                with right :
                    with st.spinner(f"Converting into Word...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            def get_key_by_value(dictionary, value) :
                                for key, val in dictionary.items() :
                                    if val == value :
                                        return key
                                # If the value is not found, you can return None or raise an exception.
                                return None

                            st.success(f"File converted to {selected_option} successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.{get_key_by_value(options, selected_option)}"

                            st.download_button("Download Converted File", data=output_file, file_name=new_file_name)
                            # todo-fixme done




    import docx2txt

    def docx_to_html(uploaded_file) :
        # Create a BytesIO object to store the converted HTML
        output_html = io.BytesIO()

        # Read the DOCX file
        docx_content = uploaded_file.read()

        # Convert DOCX to HTML
        html_content = docx2txt.process(BytesIO(docx_content))

        # Write the HTML content to the BytesIO object
        output_html.write(html_content.encode())
        output_html.seek(0)

        left, center, right = st.columns(3)
        with left :
            if st.button("Convert") :
                with right :
                    with st.spinner("Converting to HTML...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            st.success("File converted to HTML successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.html"

                            st.download_button("Download Converted File", data=output_html, file_name=new_file_name)

                            # todo-fixme done

    def csv_to_pdf(uploaded_file) :
        # Open the uploaded CSV file in text mode
        csv_file = io.TextIOWrapper(uploaded_file)

        # Read the CSV data
        csv_data = []
        reader = csv.reader(csv_file)
        for row in reader :
            csv_data.append(row)

        # Set up the PDF document
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Set font and size
        pdf.set_font("Arial", size=10)

        # Add table content
        for row in csv_data :
            for item in row :
                pdf.cell(40, 10, txt=str(item), border=1)
            pdf.ln()

        # Create a temporary file to save the PDF
        with tempfile.NamedTemporaryFile(delete=False) as temp_file :
            temp_file_path = temp_file.name
            pdf.output(temp_file_path)

        # Read the temporary PDF file as bytes
        with open(temp_file_path, 'rb') as temp_file :
            pdf_bytes = temp_file.read()

        # Display success message and download button
        st.success("File converted to PDF successfully.")
        st.download_button("Download Converted File", data=pdf_bytes, file_name='converted.pdf')

    def csv_to_text(uploaded_file) :
        # Open the uploaded CSV file in text mode
        csv_file = io.TextIOWrapper(uploaded_file)

        # Read the CSV data
        csv_data = []
        reader = csv.reader(csv_file)
        for row in reader :
            csv_data.append(row)

        # Create a StringIO object to store the converted text
        output_text = io.StringIO()

        # Convert CSV data to text
        for row in csv_data :
            line = ','.join(row)
            output_text.write(line + '\n')

        # Reset the pointer of the StringIO object
        output_text.seek(0)

        left, center, right = st.columns(3)
        with left :

            if st.button("Convert") :
                with right :
                    with st.spinner(f"Converting into Word...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            def get_key_by_value(dictionary, value) :
                                for key, val in dictionary.items() :
                                    if val == value :
                                        return key
                                # If the value is not found, you can return None or raise an exception.
                                return None

                            st.success(f"File converted to {selected_option} successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.{get_key_by_value(options, selected_option)}"
                            # convert ouput_Text to bytes-like object
                            output_bytes = output_text.getvalue().encode()

                            # Create a temporary file to store the converted text
                            temp_text_file = tempfile.NamedTemporaryFile(delete=False)
                            temp_text_file.write(output_bytes)

                            st.download_button("Download Converted File", data=temp_text_file.name,
                                               file_name=new_file_name)

                            # Close and remove the temporary file
                            temp_text_file.close()
                            os.unlink(temp_text_file.name)

    def csv_to_word(uploaded_file) :
        # Open the uploaded CSV file in text mode
        csv_file = io.TextIOWrapper(uploaded_file)

        # Read the CSV data
        csv_data = []
        reader = csv.reader(csv_file)
        for row in reader :
            csv_data.append(row)

        # Create a new Word document
        document = Document()

        # Add table content
        table = document.add_table(rows=len(csv_data), cols=max(len(row) for row in csv_data))
        for i, row in enumerate(csv_data) :
            for j, cell_value in enumerate(row) :
                table.cell(i, j).text = cell_value

        # Create a BytesIO object to save the Word document
        output_file = io.BytesIO()
        document.save(output_file)

        # Reset the file pointer of the BytesIO object
        output_file.seek(0)

        left, center, right = st.columns(3)
        with left :

            if st.button("Convert") :
                with right :
                    with st.spinner(f"Converting into Word...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            def get_key_by_value(dictionary, value) :
                                for key, val in dictionary.items() :
                                    if val == value :
                                        return key
                                # If the value is not found, you can return None or raise an exception.
                                return None

                            st.success(f"File converted to {selected_option} successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.{get_key_by_value(options, selected_option)}"

                            st.download_button("Download Converted File", data=output_file, file_name=new_file_name)

    def csv_to_html(uploaded_file) :
        # Read the CSV file
        csv_data = []
        with io.TextIOWrapper(uploaded_file) as csv_file :
            reader = csv.reader(csv_file)
            for row in reader :
                csv_data.append(row)

        # Create an HTML table string
        html_table = "<table>"
        for row in csv_data :
            html_table += "<tr>"
            for cell in row :
                html_table += f"<td>{cell}</td>"
            html_table += "</tr>"
        html_table += "</table>"

        # Create a BytesIO object to store the converted HTML
        output_html = io.BytesIO()
        output_html.write(html_table.encode())
        output_html.seek(0)

        left, center, right = st.columns(3)
        with left :
            if st.button("Convert") :
                with right :
                    with st.spinner("Converting to HTML...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            st.success("File converted to HTML successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.html"

                            st.download_button("Download Converted File", data=output_html, file_name=new_file_name)

                            # todo-fixme done

    # csv to other ends---------------------

    # html to others begins----------------------
    from docx import Document
    from fpdf import FPDF

    def html_to_docx(uploaded_file) :
        # Read the HTML content
        html_content = uploaded_file.read().decode("utf-8")

        # Create a new Word document
        document = Document()

        # Add the HTML content to the document
        document.add_paragraph(html_content)

        # Create a BytesIO object to save the Word document
        output_file = io.BytesIO()
        document.save(output_file)

        # Reset the file pointer of the BytesIO object
        output_file.seek(0)

        left, center, right = st.columns(3)
        with left :

            if st.button("Convert") :
                with right :
                    with st.spinner(f"Converting into Word...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            def get_key_by_value(dictionary, value) :
                                for key, val in dictionary.items() :
                                    if val == value :
                                        return key
                                # If the value is not found, you can return None or raise an exception.
                                return None

                            st.success(f"File converted to {selected_option} successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.{get_key_by_value(options, selected_option)}"

                            st.download_button("Download Converted File", data=output_file, file_name=new_file_name)

    import csv

    def html_to_csv(uploaded_file) :
        # Read the HTML content

        st.error("I apologize for the inconvenience. This feature will be updated later. "
                 "In the meantime, please try different formats or consider converting HTML to PDF "
                 "and then PDF to CSV. Thank you for your understanding.")

    import io
    import tempfile
    from bs4 import BeautifulSoup

    import streamlit as st

    def html_to_text(uploaded_file) :
        # Read the HTML content
        html_content = uploaded_file.read().decode("utf-8")

        # Parse the HTML using BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')

        # Extract text from HTML
        text = soup.get_text(strip=True)

        # Create a StringIO object to store the converted text
        output_text = io.StringIO()
        output_text.write(text)

        # Reset the pointer of the StringIO object
        output_text.seek(0)

        # Convert StringIO to bytes
        output_bytes = output_text.getvalue().encode()

        left, center, right = st.columns(3)
        with left :

            if st.button("Convert") :
                with right :
                    with st.spinner(f"Converting into Word...") :
                        time.sleep(3)  # Simulating conversion process
                        with left :
                            def get_key_by_value(dictionary, value) :
                                for key, val in dictionary.items() :
                                    if val == value :
                                        return key
                                # If the value is not found, you can return None or raise an exception.
                                return None

                            st.success(f"File converted to {selected_option} successfully.")

                            file_extension = uploaded_file.name.split('.')[-1]
                            new_file_name = f"{uploaded_file.name[:-len(file_extension) - 1]}.{get_key_by_value(options, selected_option)}"

                            st.download_button("Download Converted File", data=output_bytes, file_name=new_file_name)

    # html to others ends--------------------------
    # conversions from pdf to other formats starts-----------

    # conversions from pdf to other format ends

    st.markdown(
        """
        <link rel="stylesheet" href="./styles.css">
        """,
        unsafe_allow_html=True
    )

    uploaded_file = st.file_uploader("Upload a file")
    if uploaded_file is not None:

        if uploaded_file.name.split(".")[-1].lower() == 'pdf' :
            st.warning(
                "Apologies, but the PDF converter is currently unavailable. We are unable to convert PDF files at the moment. Please try again later.")
        else :

                file_extension = uploaded_file.name.split(".")[-1].lower()

                options = {
                    'docx' : 'Word',  # done
                    'csv' : 'CSV',  # done
                    'pdf' : 'PDF',  # done
                    'txt' : 'TXT',  # done
                    'html' : 'HTML',
                }

                # Remove uploaded file extension from options
                if file_extension in options :
                    options.pop(file_extension)

                # Remove HTML option if uploaded file is DOCX
                # if file_extension == 'docx' :
                #     options.pop('html')

                selected_option = st.selectbox("Convert into:",
                                               options=list(options.values()),
                                               key="option_menu",
                                               index=0,
                                               help="Select a conversion option")

                if file_extension == 'txt' :
                    if selected_option == 'HTML' :
                        text_to_html(uploaded_file)
                    if selected_option == 'Word' :
                        text_to_word(uploaded_file)
                    elif selected_option == 'CSV' :
                        text_to_csv(uploaded_file)
                    elif selected_option == 'PDF' :
                        text_to_pdf(uploaded_file)

                # Add other conversion options
                if file_extension == 'docx' :
                    if selected_option == 'HTML' :
                        docx_to_html(uploaded_file)
                    if selected_option == 'TXT' :
                        docx_to_text(uploaded_file)
                    if selected_option == 'CSV' :
                        docx_to_csv(uploaded_file)
                    if selected_option == 'PDF' :
                        docx_to_pdf(uploaded_file)
                if file_extension == 'csv' :
                    if selected_option == 'HTML' :
                        csv_to_html(uploaded_file)
                    if selected_option == 'TXT' :
                        csv_to_text(uploaded_file)
                    if selected_option == 'Word' :
                        csv_to_word(uploaded_file)
                    if selected_option == 'PDF' :
                        csv_to_pdf(uploaded_file)
                if file_extension == 'html' :

                    if selected_option == 'CSV' :
                        html_to_csv(uploaded_file)
                    if selected_option == 'Word' :
                        html_to_docx(uploaded_file)
                    if selected_option == 'TXT' :
                        html_to_text(uploaded_file)
                    if selected_option == 'PDF':
                        st.warning("I apologize for the inconvenience. Unfortunately, the PDF conversion feature is currently unavailable. We are actively working on updating it. Please try again later.")



# All site operations ends here

# About page content----------------------------------------------------
if selected == 'Home' :
    home_page()
elif selected == 'About' :
    about_page()
elif selected == 'Help/FAQ' :
    help_faq_page()
elif selected == 'Contact' :
    contact_page()
elif selected == 'Feedback' :
    feedback_page()
